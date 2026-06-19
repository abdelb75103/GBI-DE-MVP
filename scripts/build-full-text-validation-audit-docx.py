#!/usr/bin/env python3
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/research/original-search-full-text-review-skill-validation-final-audit-2026-06-19.md"
OUTPUT = ROOT / "outputs/full-text-skill-validation/original-search-full-text-review-skill-validation-final-audit-2026-06-19.docx"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def format_inline(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        else:
            paragraph.add_run(part.replace("  ", " "))


def add_markdown_table(doc, rows):
    headers = [x.strip() for x in rows[0].strip().strip("|").split("|")]
    data = [[x.strip() for x in row.strip().strip("|").split("|")] for row in rows[2:]]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = [9360 // len(headers)] * len(headers)
    widths[-1] += 9360 - sum(widths)
    for i, header in enumerate(headers):
        p = table.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for values in data:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            format_inline(p, value)
            for run in p.runs:
                run.font.size = Pt(8.5)
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = "FIFA GBI  |  Full-text skill validation audit"
    header.style = styles["Normal"]
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
    add_page_field(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("FULL-TEXT REVIEW SKILL VALIDATION")
    run.bold = True
    run.font.size = Pt(23)
    run.font.color.rgb = RGBColor(0, 0, 0)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Final audit of calibration, sealed holdout, extension, and denominator regression")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(55, 55, 55)
    for label, value in (
        ("Project", "FIFA Global Burden of Injury"),
        ("Search wave", "Original-search historical decisions"),
        ("Date", "19 June 2026"),
        ("Status", "Final audit — no historical records changed"),
        ("Criteria", "FIFA GBI full-text eligibility v4 (2026-06-19)"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        r.bold = True
        p.add_run(value)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(10)
    ppr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "2E74B5")
    pbdr.append(bottom)
    ppr.append(pbdr)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 1
    while i < len(lines):
        line = lines[i].rstrip()
        if not line or line == "---":
            i += 1
            continue
        if line.startswith("**Workflow stage:") or line.startswith("**Search/import wave:") or line.startswith("**Date:") or line.startswith("**Status:") or line.startswith("**Criteria version:"):
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i + 1].strip()):
            rows = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(lines[i])
                i += 1
            add_markdown_table(doc, rows)
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 3")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("# "):
            pass
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            format_inline(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            format_inline(p, line[2:])
        else:
            p = doc.add_paragraph()
            format_inline(p, line)
        i += 1

    props = doc.core_properties
    props.title = "Original-search full-text review skill validation — final audit"
    props.subject = "FIFA GBI full-text screening skill validation"
    props.author = "FIFA GBI project"
    props.keywords = "FIFA GBI, full-text screening, validation, audit"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
